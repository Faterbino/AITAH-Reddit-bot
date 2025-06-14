import praw
import re
import docx
from pathlib import Path
from datetime import date
from update_remover import Update_Remover

reddit = praw.Reddit("bot1")

docx_path = Path("AITAH stories.docx")

if docx_path.is_file():
    doc=docx.Document(docx_path)
    doc.add_page_break()
    doc.add_heading(str(date.today()),0)
else:
    doc = docx.Document()
    doc.save(docx_path)
    doc.add_heading(str(date.today()),0) 

subreddit = reddit.subreddit("AITAH")
for submission in subreddit.hot(limit=10):
    if re.search("update",submission.title,re.IGNORECASE):
        title = Update_Remover(submission.title)
        user = reddit.redditor(f"{submission.author}")
        for user_submission in user.submissions.new(limit=None):
            if re.search(title,user_submission.title,re.IGNORECASE) and not re.search("update",user_submission.title,re.IGNORECASE):
                doc.add_heading(user_submission.title,1)
                doc.add_paragraph(user_submission.selftext)          
    doc.add_heading(submission.title,1)
    doc.add_paragraph(submission.selftext)

doc.save(docx_path)